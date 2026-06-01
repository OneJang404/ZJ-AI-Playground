"""
DOCX 文本与结构提取模块
======================
用 python-docx 提取 Word 文档的文本、表格、章节结构，
对标 PyMuPDF 的文字层提取能力，不受印章/图片遮挡影响。
"""
import io
import re
import logging
import subprocess
import tempfile
import os
from pathlib import Path
from typing import Tuple, Optional

logger = logging.getLogger(__name__)

# ---- LibreOffice 可执行文件路径（Windows / macOS / Linux） ----
_LIBREOFFICE_PATHS = [
    "soffice",                         # PATH 中
    "libreoffice",
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
    "/usr/bin/soffice",                                      # Linux
]


def _find_libreoffice() -> Optional[str]:
    """查找 LibreOffice 可执行文件，找不到返回 None"""
    for p in _LIBREOFFICE_PATHS:
        if os.path.exists(p):
            return p
    # 尝试 which
    try:
        result = subprocess.run(
            ["where", "soffice"], capture_output=True, text=True, shell=True, timeout=5
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip().split("\n")[0]
    except Exception:
        pass
    return None


def convert_docx_to_pdf(docx_bytes: bytes) -> Optional[bytes]:
    """
    用 LibreOffice headless 将 DOCX 转为 PDF 字节流
    返回 PDF 字节，失败返回 None
    """
    lo_path = _find_libreoffice()
    if not lo_path:
        logger.warning("未找到 LibreOffice，无法将 DOCX 渲染为 PDF。"
                       "请安装：https://www.libreoffice.org/download/")
        return None

    with tempfile.TemporaryDirectory(prefix="docx2pdf_") as tmpdir:
        docx_path = os.path.join(tmpdir, "input.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        try:
            result = subprocess.run(
                [lo_path, "--headless", "--convert-to", "pdf",
                 "--outdir", tmpdir, docx_path],
                capture_output=True, text=True, timeout=120,
            )
            pdf_path = os.path.join(tmpdir, "input.pdf")
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    return f.read()
            else:
                logger.warning(f"LibreOffice 转换失败：{result.stderr[:300]}")
                return None
        except subprocess.TimeoutExpired:
            logger.warning("LibreOffice 转换超时（>120s），DOCX 可能过大或含有复杂元素")
            return None
        except Exception as e:
            logger.warning(f"LibreOffice 转换异常：{e}")
            return None


def extract_docx_text(docx_bytes: bytes) -> Tuple[str, int]:
    """
    从 DOCX 字节流提取所有文本（含表格），对标 _extract_text_only( PDF )
    返回: (格式化文本, 估算页数)
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(docx_bytes))
    parts = []
    _extract_blocks(doc.element.body, doc, parts)

    full_text = "\n".join(parts)
    # 估算页数：中文文档约 800 字符/页
    est_pages = max(1, len(full_text) // 800)
    logger.info(f"DOCX 文本提取：{len(full_text)} 字符，估计 {est_pages} 页")
    return full_text, est_pages


def _extract_blocks(body, doc, parts: list):
    """递归提取文档主体中的段落和表格"""
    from docx.oxml.ns import qn

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            text = _get_paragraph_text(child, doc)
            if text.strip():
                parts.append(text.strip())
        elif tag == "tbl":
            table_text = _extract_table_text(child, doc)
            if table_text:
                parts.append(table_text)


def _get_paragraph_text(p_element, doc) -> str:
    """获取段落文本（含超链接、样式标记）"""
    from docx.oxml.ns import qn

    runs = []
    for r in p_element.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is not None and t.text:
            runs.append(t.text)
        # 检查换行符
        br = r.find(qn("w:br"))
        if br is not None:
            runs.append("\n")
    return "".join(runs)


def _extract_table_text(tbl_element, doc) -> str:
    """提取表格文本，保留行列结构"""
    from docx.oxml.ns import qn

    rows = []
    for tr in tbl_element.findall(qn("w:tr")):
        cells = []
        for tc in tr.findall(qn("w:tc")):
            cell_parts = []
            for p in tc.findall(qn("w:p")):
                text = _get_paragraph_text(p, doc)
                if text.strip():
                    cell_parts.append(text.strip())
            cells.append(" ".join(cell_parts))
        rows.append(" | ".join(cells))
    if rows:
        return "\n".join(rows)
    return ""


# ---- 便捷函数：判断文件类型 ----

# OLE2 复合文档魔数（.doc / .xls / .ppt 等旧格式）
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def is_docx(file_bytes: bytes, filename: str = "") -> bool:
    """判断文件是否为 DOCX 格式（Office Open XML）"""
    if filename.lower().endswith(".docx"):
        return True
    if file_bytes[:2] == b"PK":
        try:
            import zipfile
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                return "[Content_Types].xml" in zf.namelist()
        except Exception:
            pass
    return False


def is_doc(file_bytes: bytes, filename: str = "") -> bool:
    """判断文件是否为旧版 .doc 格式（OLE2 复合文档）"""
    if filename.lower().endswith(".doc"):
        return True
    return file_bytes[:8] == _OLE2_MAGIC


def is_word_doc(file_bytes: bytes, filename: str = "") -> bool:
    """判断是否为 Word 文档（.doc 或 .docx）"""
    return is_docx(file_bytes, filename) or is_doc(file_bytes, filename)


# 招投标文档中常见的项目名称前置标记
_PROJECT_NAME_PATTERNS = [
    r"项目名称[：:]\s*(.+)",
    r"招标项目[：:]\s*(.+)",
    r"采购项目名称[：:]\s*(.+)",
    r"工程名称[：:]\s*(.+)",
    r"项目[：:]\s*(.+)",
    r"标段名称[：:]\s*(.+)",
    r"招标项目名称[：:]\s*(.+)",
    r"采购项目[：:]\s*(.+)",
]


def extract_project_name(file_bytes: bytes, filename: str = "") -> str:
    """
    从上传文件的头几页提取招标项目名称。
    支持 PDF / DOCX，优先匹配「项目名称：XXX」等常见字段，
    匹配不到则回退到文件名。

    返回: 项目名称字符串
    """
    text = ""
    # ---- PDF ----
    if not is_word_doc(file_bytes, filename):
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages = min(3, len(doc))
            parts = []
            for i in range(pages):
                t = doc[i].get_text()
                if t.strip():
                    parts.append(t.strip())
            doc.close()
            text = "\n".join(parts)
        except Exception:
            pass
    # ---- DOCX ----
    else:
        try:
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(file_bytes))
            parts = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    parts.append(para.text.strip())
                if i > 80:  # 头几页的段落数
                    break
            text = "\n".join(parts)
        except Exception:
            pass

    if not text:
        return filename.rsplit(".", 1)[0] if filename else "未知项目"

    # 逐行匹配项目名称模式
    for line in text.split("\n"):
        line = line.strip()
        if len(line) > 50 or len(line) < 3:
            continue
        for pattern in _PROJECT_NAME_PATTERNS:
            m = re.search(pattern, line)
            if m:
                name = m.group(1).strip()
                # 清理尾随的标点
                name = name.rstrip("，。；、")
                if 2 <= len(name) <= 60:
                    return name

    # 回退：取头几页中看起来像标题的第一行非空文字
    for line in text.split("\n")[:5]:
        line = line.strip()
        if 4 <= len(line) <= 40 and not line.startswith("第") and "页" not in line:
            return line

    return filename.rsplit(".", 1)[0] if filename else "未知项目"
